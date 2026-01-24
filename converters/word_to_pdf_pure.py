"""
Word to PDF Converter - VERSIONE PURA PYTHON

NON serve Office/LibreOffice!
Usa solo librerie Python: python-docx + reportlab
"""

from pathlib import Path
import platform

from core.converter_base import ConverterBase
from utils.error_handler import ConversionError


class WordToPDFConverter(ConverterBase):
    """
    Convertitore Word → PDF usando SOLO Python
    NO Office/LibreOffice necessario!
    """
    
    def __init__(self):
        """Inizializza il convertitore Word → PDF"""
        super().__init__()
        self._system = platform.system().lower()
    
    def get_info(self) -> dict:
        """Informazioni sul convertitore"""
        return {
            'name': 'Word to PDF',
            'description': 'Converte documenti Word (.docx, .doc) in PDF',
            'input_formats': ['.docx', '.doc'],
            'output_format': '.pdf',
            'requires_dep': None  # ✅ NO dipendenze esterne!
        }
    
    def convert(self, input_file: Path, output_file: Path, progress_callback=None) -> bool:
        """
        Converte Word → PDF usando SOLO librerie Python
        
        Args:
            input_file: File Word di input
            output_file: File PDF di output
            progress_callback: Callback per progress bar
            
        Returns:
            True se successo
        """
        try:
            self.logger.info(f"Inizio conversione: {input_file} → {output_file}")
            self._report_progress(progress_callback, 10, "Lettura documento Word...")
            
            # ✅ METODO PURO PYTHON
            return self._convert_with_python_libs(input_file, output_file, progress_callback)
            
        except Exception as e:
            self.logger.error(f"Errore conversione: {e}")
            raise ConversionError(
                f"Errore conversione Word→PDF: {str(e)}",
                file_path=str(input_file)
            )
    
    def _convert_with_python_libs(self, input_file: Path, output_file: Path, progress_callback=None) -> bool:
        """
        ⚡ CONVERSIONE PURA PYTHON - NO Office necessario!
        
        Usa:
        - python-docx: Legge .docx
        - reportlab: Crea PDF
        """
        try:
            from docx import Document
            from reportlab.lib.pagesizes import letter, A4
            from reportlab.lib.units import inch
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_RIGHT, TA_JUSTIFY
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
            from reportlab.lib import colors
            
            self._report_progress(progress_callback, 20, "Lettura documento Word...")
            
            # Leggi documento Word
            doc = Document(str(input_file))
            
            self._report_progress(progress_callback, 40, "Creazione PDF...")
            
            # Crea PDF con ReportLab
            pdf_doc = SimpleDocTemplate(
                str(output_file),
                pagesize=A4,
                rightMargin=72,
                leftMargin=72,
                topMargin=72,
                bottomMargin=18,
            )
            
            # Stili
            styles = getSampleStyleSheet()
            story = []
            
            self._report_progress(progress_callback, 60, "Conversione contenuto...")
            
            # Converti paragrafi
            for para in doc.paragraphs:
                if para.text.strip():
                    # Determina stile
                    if para.style.name.startswith('Heading'):
                        style = styles['Heading1']
                    else:
                        style = styles['Normal']
                    
                    # Aggiungi paragrafo
                    p = Paragraph(para.text, style)
                    story.append(p)
                    story.append(Spacer(1, 0.2*inch))
            
            # Converti tabelle (semplificato)
            for table in doc.tables:
                for row in table.rows:
                    row_text = ' | '.join([cell.text for cell in row.cells])
                    if row_text.strip():
                        p = Paragraph(row_text, styles['Normal'])
                        story.append(p)
                        story.append(Spacer(1, 0.1*inch))
                
                story.append(Spacer(1, 0.3*inch))
            
            self._report_progress(progress_callback, 80, "Salvataggio PDF...")
            
            # Genera PDF
            pdf_doc.build(story)
            
            self._report_progress(progress_callback, 100, "Completato!")
            
            if output_file.exists() and output_file.stat().st_size > 0:
                self.logger.info(f"✅ Conversione completata: {output_file}")
                return True
            
            return False
            
        except ImportError as e:
            self.logger.error(f"Librerie mancanti: {e}")
            raise ConversionError(
                f"Librerie Python mancanti!\n\n"
                f"Installa con:\n"
                f"pip install python-docx reportlab\n\n"
                f"Errore: {e}"
            )
        
        except Exception as e:
            self.logger.error(f"Errore conversione Python: {e}")
            raise ConversionError(f"Errore conversione: {str(e)}")

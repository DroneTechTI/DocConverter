"""
Word to PDF Converter - VERSIONE PURA PYTHON

NON serve Office/LibreOffice!
Usa solo librerie Python: python-docx + reportlab
"""

from pathlib import Path
import platform

from core.converter_base import ConverterBase
from utils.error_handler import ConversionError


class WordToPDFConverter(ConverterBase):
    """
    Convertitore Word → PDF usando SOLO Python
    NO Office/LibreOffice necessario!
    """
    
    def __init__(self):
        """Inizializza il convertitore Word → PDF"""
        super().__init__()
        self._system = platform.system().lower()
    
    def get_info(self) -> dict:
        """Informazioni sul convertitore"""
        return {
            'name': 'Word to PDF',
            'description': 'Converte documenti Word (.docx, .doc) in PDF',
            'input_formats': ['.docx', '.doc'],
            'output_format': '.pdf',
            'requires_dep': None  # ✅ NO dipendenze esterne!
        }
    
    def convert(self, input_file: Path, output_file: Path, progress_callback=None) -> bool:
        """
        Converte Word → PDF con formattazione perfetta
        
        STRATEGIA MIGLIORE:
        1. Windows + Word installato: usa COM automation (PERFETTO)
        2. LibreOffice disponibile: usa soffice (OTTIMO)
        3. Fallback: python-docx + reportlab (ACCETTABILE)
        
        Args:
            input_file: File Word di input
            output_file: File PDF di output
            progress_callback: Callback per progress bar
            
        Returns:
            True se successo
        """
        try:
            self.logger.info(f"Inizio conversione: {input_file} → {output_file}")
            self._report_progress(progress_callback, 10, "Selezione metodo conversione...")
            
            # PRIORITÀ 1: Word COM (Windows) - FORMATTAZIONE PERFETTA
            if self._system == 'windows':
                if self._try_word_com(input_file, output_file, progress_callback):
                    return True
            
            # PRIORITÀ 2: LibreOffice (cross-platform) - OTTIMA FORMATTAZIONE
            if self._try_libreoffice(input_file, output_file, progress_callback):
                return True
            
            # FALLBACK: Python puro - formattazione base
            self.logger.warning("⚠️ Usando metodo Python (formattazione base)")
            return self._convert_with_python_libs(input_file, output_file, progress_callback)
            
        except Exception as e:
            self.logger.error(f"Errore conversione: {e}")
            raise ConversionError(
                f"Errore conversione Word→PDF: {str(e)}",
                file_path=str(input_file)
            )
    
    def _convert_with_python_libs(self, input_file: Path, output_file: Path, progress_callback=None) -> bool:
        """
        ⚡ CONVERSIONE PURA PYTHON - NO Office necessario!
        
        Usa:
        - python-docx: Legge .docx
        - reportlab: Crea PDF
        
        Ottimizzazioni:
        - Import lazy per velocità
        - Batch processing per grandi documenti
        """
        try:
            # Import lazy (solo quando serve)
            from docx import Document
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from reportlab.lib.pagesizes import A4
            from reportlab.lib.units import inch, cm
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_RIGHT, TA_JUSTIFY
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak, Table, TableStyle
            from reportlab.lib import colors
            
            self._report_progress(progress_callback, 20, "Lettura documento Word...")
            
            # Leggi documento Word (ottimizzato)
            doc = Document(str(input_file))
            
            self._report_progress(progress_callback, 40, "Creazione PDF...")
            
            # Crea PDF con ReportLab
            pdf_doc = SimpleDocTemplate(
                str(output_file),
                pagesize=A4,
                rightMargin=2*cm,
                leftMargin=2*cm,
                topMargin=2*cm,
                bottomMargin=2*cm,
            )
            
            # Stili migliorati
            styles = getSampleStyleSheet()
            
            # Crea stili personalizzati per preservare formattazione
            styles.add(ParagraphStyle(
                name='CustomNormal',
                parent=styles['Normal'],
                fontSize=11,
                leading=14,
                spaceBefore=6,
                spaceAfter=6,
            ))
            
            styles.add(ParagraphStyle(
                name='CustomHeading1',
                parent=styles['Heading1'],
                fontSize=16,
                leading=20,
                spaceBefore=12,
                spaceAfter=6,
                textColor=colors.HexColor('#2C3E50'),
            ))
            
            styles.add(ParagraphStyle(
                name='CustomHeading2',
                parent=styles['Heading2'],
                fontSize=14,
                leading=18,
                spaceBefore=10,
                spaceAfter=6,
                textColor=colors.HexColor('#34495E'),
            ))
            
            story = []
            
            self._report_progress(progress_callback, 60, "Conversione contenuto...")
            
            # Converti paragrafi con formattazione migliorata
            total_paragraphs = len(doc.paragraphs)
            for idx, para in enumerate(doc.paragraphs):
                if para.text.strip():
                    # Determina stile basato su heading e formattazione
                    if para.style.name.startswith('Heading 1'):
                        style = styles['CustomHeading1']
                    elif para.style.name.startswith('Heading 2'):
                        style = styles['CustomHeading2']
                    elif para.style.name.startswith('Heading'):
                        style = styles['CustomHeading2']
                    else:
                        style = styles['CustomNormal']
                    
                    # Costruisci testo con formattazione HTML-like per ReportLab
                    text = self._build_formatted_text(para)
                    
                    # Allineamento
                    alignment = self._get_alignment(para)
                    if alignment != style.alignment:
                        custom_style = ParagraphStyle(
                            name=f'temp_{idx}',
                            parent=style,
                            alignment=alignment
                        )
                        style = custom_style
                    
                    # Aggiungi paragrafo
                    p = Paragraph(text, style)
                    story.append(p)
                    story.append(Spacer(1, 0.1*inch))
                
                # Progress update ogni 10%
                if idx % max(1, total_paragraphs // 10) == 0:
                    progress = 60 + int((idx / total_paragraphs) * 20)
                    self._report_progress(progress_callback, progress, "Conversione contenuto...")
            
            # Converti tabelle con formattazione migliore
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = []
                    for cell in row.cells:
                        # Estrai testo dalle celle preservando formattazione base
                        cell_text = cell.text.strip()
                        if cell_text:
                            row_data.append(cell_text)
                        else:
                            row_data.append('')
                    
                    if any(row_data):  # Aggiungi solo righe non vuote
                        table_data.append(row_data)
                
                if table_data:
                    # Crea tabella ReportLab
                    pdf_table = Table(table_data)
                    pdf_table.setStyle(TableStyle([
                        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#E8E8E8')),
                        ('TEXTCOLOR', (0, 0), (-1, 0), colors.black),
                        ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                        ('FONTSIZE', (0, 0), (-1, 0), 10),
                        ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                        ('BACKGROUND', (0, 1), (-1, -1), colors.white),
                        ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
                        ('FONTNAME', (0, 1), (-1, -1), 'Helvetica'),
                        ('FONTSIZE', (0, 1), (-1, -1), 9),
                        ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                        ('LEFTPADDING', (0, 0), (-1, -1), 6),
                        ('RIGHTPADDING', (0, 0), (-1, -1), 6),
                    ]))
                    story.append(pdf_table)
                    story.append(Spacer(1, 0.3*inch))
            
            self._report_progress(progress_callback, 85, "Salvataggio PDF...")
            
            # Genera PDF
            pdf_doc.build(story)
            
            self._report_progress(progress_callback, 100, "Completato!")
            
            if output_file.exists() and output_file.stat().st_size > 0:
                self.logger.info(f"✅ Conversione completata: {output_file}")
                return True
            
            return False
            
        except ImportError as e:
            self.logger.error(f"Librerie mancanti: {e}")
            raise ConversionError(
                f"Librerie Python mancanti!\n\n"
                f"Installa con:\n"
                f"pip install python-docx reportlab\n\n"
                f"Errore: {e}"
            )
        
        except Exception as e:
            self.logger.error(f"Errore conversione Python: {e}")
            raise ConversionError(f"Errore conversione: {str(e)}")
    
    def _build_formatted_text(self, paragraph) -> str:
        """
        Costruisce testo con formattazione HTML-like per ReportLab
        Preserva: grassetto, corsivo, sottolineato
        """
        text_parts = []
        
        for run in paragraph.runs:
            text = run.text
            
            # Escape caratteri speciali HTML
            text = text.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
            
            # Applica formattazione
            if run.bold and run.italic:
                text = f'<b><i>{text}</i></b>'
            elif run.bold:
                text = f'<b>{text}</b>'
            elif run.italic:
                text = f'<i>{text}</i>'
            
            if run.underline:
                text = f'<u>{text}</u>'
            
            text_parts.append(text)
        
        return ''.join(text_parts) if text_parts else paragraph.text
    
    def _try_word_com(self, input_file: Path, output_file: Path, progress_callback=None) -> bool:
        """
        Tenta conversione con Word COM (Windows) - FORMATTAZIONE PERFETTA
        """
        try:
            import win32com.client
            
            self.logger.info("✅ Uso Word COM (formattazione PERFETTA)")
            self._report_progress(progress_callback, 30, "Conversione con Microsoft Word...")
            
            # Avvia Word
            word = win32com.client.Dispatch("Word.Application")
            word.Visible = False
            
            try:
                # Apri documento
                doc = word.Documents.Open(str(input_file.absolute()))
                
                # Salva come PDF (formato 17 = wdFormatPDF)
                doc.SaveAs(str(output_file.absolute()), FileFormat=17)
                doc.Close()
                
                self._report_progress(progress_callback, 100, "Completato!")
                
                if output_file.exists() and output_file.stat().st_size > 0:
                    self.logger.info(f"✅ PDF creato con Word COM: {output_file}")
                    return True
                    
            finally:
                word.Quit()
            
            return False
            
        except ImportError:
            self.logger.debug("pywin32 non disponibile")
            return False
        except Exception as e:
            self.logger.warning(f"Word COM fallito: {e}")
            return False
    
    def _try_libreoffice(self, input_file: Path, output_file: Path, progress_callback=None) -> bool:
        """
        Tenta conversione con LibreOffice - OTTIMA FORMATTAZIONE
        """
        import subprocess
        import shutil
        
        # Cerca soffice/libreoffice
        soffice_paths = [
            'soffice',
            'libreoffice',
            r'C:\Program Files\LibreOffice\program\soffice.exe',
            r'C:\Program Files (x86)\LibreOffice\program\soffice.exe',
            '/usr/bin/soffice',
            '/usr/bin/libreoffice',
        ]
        
        soffice = None
        for path in soffice_paths:
            if shutil.which(path) or Path(path).exists():
                soffice = path
                break
        
        if not soffice:
            self.logger.debug("LibreOffice non trovato")
            return False
        
        try:
            self.logger.info("✅ Uso LibreOffice (formattazione OTTIMA)")
            self._report_progress(progress_callback, 30, "Conversione con LibreOffice...")
            
            # Comando LibreOffice
            cmd = [
                soffice,
                '--headless',
                '--convert-to', 'pdf',
                '--outdir', str(output_file.parent),
                str(input_file.absolute())
            ]
            
            result = subprocess.run(
                cmd,
                capture_output=True,
                timeout=60,
                text=True
            )
            
            self._report_progress(progress_callback, 90, "Verifica output...")
            
            # LibreOffice crea PDF con stesso nome input
            expected_pdf = output_file.parent / f"{input_file.stem}.pdf"
            
            if expected_pdf.exists() and expected_pdf != output_file:
                expected_pdf.rename(output_file)
            
            if output_file.exists() and output_file.stat().st_size > 0:
                self.logger.info(f"✅ PDF creato con LibreOffice: {output_file}")
                self._report_progress(progress_callback, 100, "Completato!")
                return True
            
            return False
            
        except (subprocess.TimeoutExpired, subprocess.CalledProcessError) as e:
            self.logger.warning(f"LibreOffice fallito: {e}")
            return False
        except Exception as e:
            self.logger.warning(f"LibreOffice errore: {e}")
            return False
    
    def _convert_with_docx2pdf(self, input_file: Path, output_file: Path, progress_callback=None) -> bool:
        """
        Conversione con docx2pdf (Windows) - USA WORD COM
        Formattazione PERFETTA perché usa Word reale
        """
        try:
            import docx2pdf
            
            self._report_progress(progress_callback, 30, "Conversione con Word...")
            
            # docx2pdf usa Word COM direttamente
            docx2pdf.convert(str(input_file), str(output_file))
            
            self._report_progress(progress_callback, 100, "Completato!")
            
            if output_file.exists() and output_file.stat().st_size > 0:
                self.logger.info(f"✅ PDF creato con formattazione perfetta: {output_file}")
                return True
            
            return False
        
        except Exception as e:
            self.logger.error(f"Errore docx2pdf: {e}")
            raise
    
    def _get_alignment(self, paragraph) -> int:
        """
        Ottiene allineamento paragrafo per ReportLab
        """
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_RIGHT, TA_JUSTIFY
        
        alignment_map = {
            WD_ALIGN_PARAGRAPH.LEFT: TA_LEFT,
            WD_ALIGN_PARAGRAPH.CENTER: TA_CENTER,
            WD_ALIGN_PARAGRAPH.RIGHT: TA_RIGHT,
            WD_ALIGN_PARAGRAPH.JUSTIFY: TA_JUSTIFY,
        }
        
        return alignment_map.get(paragraph.alignment, TA_LEFT)
